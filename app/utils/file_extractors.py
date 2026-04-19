"""Resume file extraction: PDF, DOCX, plain text."""
from __future__ import annotations

from pathlib import Path

from app.exceptions import ResumeExtractionError
from app.logger import get_logger


log = get_logger(__name__)


def extract_text(path: str | Path) -> str:
    p = Path(path)
    if not p.exists():
        raise ResumeExtractionError(f"Resume file not found: {p}")
    suffix = p.suffix.lower()
    log.info("Extracting text from %s (suffix=%s)", p, suffix)

    try:
        if suffix == ".pdf":
            return _extract_pdf(p)
        if suffix == ".docx":
            return _extract_docx(p)
        if suffix in {".txt", ".md", ""}:
            return p.read_text(encoding="utf-8", errors="replace")
        raise ResumeExtractionError(f"Unsupported resume format: {suffix}")
    except ResumeExtractionError:
        raise
    except Exception as e:
        raise ResumeExtractionError(f"Failed to extract text from {p}: {e}") from e


def _extract_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ResumeExtractionError("pypdf is not installed") from e

    reader = PdfReader(str(p))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:
            log.warning("Failed to extract a PDF page: %s", e)
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeExtractionError(f"PDF produced no extractable text: {p}")
    return text


def _extract_docx(p: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ResumeExtractionError("python-docx is not installed") from e

    doc = Document(str(p))
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(t for t in parts if t).strip()
    if not text:
        raise ResumeExtractionError(f"DOCX produced no extractable text: {p}")
    return text
